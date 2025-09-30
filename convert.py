import sys
import os
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.shared import Inches

def convert_pdf_to_docx(pdf_path, docx_path):
    # Create a new Document
    doc = Document()

    # Convert PDF to images
    images = convert_from_path(pdf_path)

    for i, image in enumerate(images):
        # Save image
        image_path = f"page_{i}.png"
        image.save(image_path, 'PNG')

        # Add image to DOCX
        doc.add_picture(image_path, width=Inches(6))

        # Extract text from image using OCR
        text = pytesseract.image_to_string(image)

        # Add text to DOCX
        doc.add_paragraph(text)

        # Remove the image file
        os.remove(image_path)

    # Save the DOCX file
    doc.save(docx_path)

if __name__ == "__main__":
    pdf_path = sys.argv[1]
    docx_path = sys.argv[2]
    convert_pdf_to_docx(pdf_path, docx_path)
